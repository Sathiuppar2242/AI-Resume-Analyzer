import os

from pypdf import PdfReader
from docx import Document


def clean_extracted_text(text):
    """Clean unnecessary whitespace from extracted resume text."""

    lines = []

    for line in text.splitlines():
        line = line.strip()

        if line:
            lines.append(line)

    return "\n".join(lines)


def extract_text_from_pdf(file_path):
    """Extract readable text from a PDF resume."""

    text_parts = []

    reader = PdfReader(file_path)

    for page_number, page in enumerate(reader.pages, start=1):

        try:
            page_text = page.extract_text()

            if page_text and page_text.strip():
                text_parts.append(page_text.strip())

        except Exception as error:
            print(
                f"PDF page {page_number} extraction error: "
                f"{type(error).__name__}: {str(error)}"
            )

            continue

    return clean_extracted_text(
        "\n".join(text_parts)
    )


def extract_text_from_docx(file_path):
    """Extract text from paragraphs and tables in a DOCX resume."""

    document = Document(file_path)

    text_parts = []

    # Extract paragraph text
    for paragraph in document.paragraphs:

        if paragraph.text.strip():
            text_parts.append(
                paragraph.text.strip()
            )

    # Extract text from tables
    for table in document.tables:

        for row in table.rows:

            for cell in row.cells:

                cell_text = cell.text.strip()

                if cell_text:
                    text_parts.append(
                        cell_text
                    )

    return clean_extracted_text(
        "\n".join(text_parts)
    )


def extract_resume_text(file_path):
    """Extract text from a supported resume file."""

    extension = os.path.splitext(
        file_path
    )[1].lower()

    if extension == ".pdf":

        return extract_text_from_pdf(
            file_path
        )

    if extension == ".docx":

        return extract_text_from_docx(
            file_path
        )

    raise ValueError(
        "Unsupported file format"
    )